import os
import json
from io import BytesIO

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------------------------------------------------------
# Configuration & helpers
# ------------------------------------------------------------------
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("OPENAI_API_KEY environment variable not found.")
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)

# Prompt copied from Next.js route (shortened where comments removed)
DOCUMENTATION_TEMPLATE = """
You are provided with a Python script. Your task is to return extremely detailed documentation in a SINGLE JSON object (no additional text). The JSON MUST follow the exact structure below and every field must be present.

Note on "tableGrain": specify WHICH columns guarantee that the final output table will contain exactly ONE row per combination of those columns.

JSON FORMAT (copy exactly – populate all placeholders):
{
  "description": "string",
  "tableGrain": "string",
  "dataSources": ["string"],
  "databricksTables": [
    { "tableName": "string", "description": "string" }
  ],
  "tableMetadata": [
    {
      "tableName": "string",
      "columns": [
        {
          "columnName": "string",
          "dataType": "string",
          "description": "string",
          "sampleValues": "string",
          "sourceTable": "string",
          "sourceColumn": "string"
        }
      ]
    }
  ],
  "integratedRules": ["string"]
}

- Populate "dataSources" with ALL input tables or files referenced in the script.
- "databricksTables" lists every table the script creates or overwrites in Databricks along with a concise business-focused description.
- "tableMetadata" must be an array, one object per output table listed in "databricksTables". Each object has tableName and columns list.
- "integratedRules" should be a BULLETED LIST (array of strings) describing transformations/business logic in order.
- For the "sourceTable" field: if the script uses a temp view/CTE, resolve to the original underlying table.
- Do NOT omit any property. Use "N/A" if genuinely unknown.
- The response MUST be valid JSON – no markdown.
"""

SYSTEM_ROLE = (
    "You are a technical documentation expert specializing in data pipeline and analytics code documentation for a business audience. "
    "Your task is to help business users understand Python code related to sales representative activities with doctors and hospitals. "
    "You create comprehensive, structured documentation following the provided template, explaining technical steps in business terms."
)


def generate_documentation(python_code: str, filename: str, stream: bool = False):
    """Call OpenAI to generate documentation JSON (optionally stream tokens)."""
    messages = [
        {"role": "system", "content": SYSTEM_ROLE},
        {
            "role": "user",
            "content": f"{DOCUMENTATION_TEMPLATE}\n\nPython file: {filename}\n\nPython Code:\n```python\n{python_code}\n```\n\nPlease generate the documentation following the exact template format provided above.",
        },
    ]

    return client.chat.completions.create(
        model="o3-2025-04-16",
        messages=messages,
        response_format={"type": "json_object"},
        stream=stream,
    )


# ------------------------------------------------------------------
# DOCX creation helpers (simple white tables)
# ------------------------------------------------------------------

def add_heading(paragraph, text, level=1):
    run = paragraph.add_run(text)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)


def build_docx(doc_data: dict, filename: str) -> BytesIO:
    doc = Document()

    # Title
    title = doc.add_heading(level=0)
    title_run = title.add_run("Python Documentation Report")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated for: {filename}")

    # 1 Description
    add_heading(doc.add_paragraph(), "1. Description", 1)
    doc.add_paragraph(doc_data.get("description", ""))

    # 2 Table Grain
    add_heading(doc.add_paragraph(), "2. Table Grain", 1)
    doc.add_paragraph(doc_data.get("tableGrain", ""))

    # 3 Data Sources
    add_heading(doc.add_paragraph(), "3. Data Sources", 1)
    for src in doc_data.get("dataSources", []):
        doc.add_paragraph(f"• {src}")

    # 4 Databricks Tables
    add_heading(doc.add_paragraph(), "4. Databricks Tables (Output)", 1)
    for t in doc_data.get("databricksTables", []):
        doc.add_paragraph(f"• {t['tableName']}: {t['description']}")

    # 5 Table Metadata
    add_heading(doc.add_paragraph(), "5. Table Metadata", 1)
    for tbl in doc_data.get("tableMetadata", []):
        add_heading(doc.add_paragraph(), f"Table: {tbl['tableName']}", 2)
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Column Name"
        hdr_cells[1].text = "Data Type"
        hdr_cells[2].text = "Description"
        hdr_cells[3].text = "Sample Values"
        hdr_cells[4].text = "Source Table"
        hdr_cells[5].text = "Source Column"
        for col in tbl.get("columns", []):
            row_cells = table.add_row().cells
            row_cells[0].text = col["columnName"]
            row_cells[1].text = col["dataType"]
            row_cells[2].text = col["description"]
            row_cells[3].text = col["sampleValues"]
            row_cells[4].text = col["sourceTable"]
            row_cells[5].text = col["sourceColumn"]
        doc.add_paragraph()

    # 6 Integrated Rules
    add_heading(doc.add_paragraph(), "6. Integrated Rules", 1)
    for rule in doc_data.get("integratedRules", []):
        doc.add_paragraph(f"• {rule}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ------------------------------------------------------------------
# Streamlit UI
# ------------------------------------------------------------------

st.set_page_config(page_title="Python Documentation Generator", layout="wide")
st.title("Python Documentation Generator (Streamlit)")

uploaded_file = st.file_uploader("Upload a Python (.py) file", type=["py"])

# Optional: choose whether to stream the response from OpenAI
stream_response = st.checkbox(
    "Stream response from OpenAI (beta – requires a supported plan)",
    value=False,
)

if uploaded_file:
    # Show immediate feedback that file was uploaded
    st.success(f"✅ File uploaded: {uploaded_file.name}")
    
    # Read and display file info
    code = uploaded_file.read().decode("utf-8")
    lines_count = len(code.split('\n'))
    st.info(f"📄 File contains {lines_count} lines of code")
    
    # Show a preview of the code
    with st.expander("Preview uploaded code (first 20 lines)"):
        preview_lines = code.split('\n')[:20]
        st.code('\n'.join(preview_lines), language='python')

    if st.button("Generate Documentation", type="primary"):
        # Initialize session state for generated docs if not exists
        if 'generated_docs' not in st.session_state:
            st.session_state.generated_docs = None
            st.session_state.generated_filename = None
        
        with st.spinner("Generating documentation…"):
            try:
                if stream_response:
                    # --- Streaming path ---
                    chunks = generate_documentation(code, uploaded_file.name, stream=True)
                    doc_string = ""

                    # Create a placeholder for streaming display
                    streaming_placeholder = st.empty()

                    for chunk in chunks:
                        delta = chunk.choices[0].delta.content if chunk.choices else ""
                        if delta:
                            doc_string += delta
                            # Update the placeholder with current progress
                            with streaming_placeholder.container():
                                st.text_area(
                                    "Generating... (live preview)",
                                    doc_string,
                                    height=200,
                                    key=f"streaming_{len(doc_string)}",
                                )

                    # Clear the streaming placeholder
                    streaming_placeholder.empty()
                else:
                    # --- Non-streaming path ---
                    response = generate_documentation(code, uploaded_file.name, stream=False)
                    doc_string = response.choices[0].message.content

                # Parse the final JSON (applies to both paths)
                doc_json = json.loads(doc_string)
                
                # Store in session state
                st.session_state.generated_docs = doc_json
                st.session_state.generated_filename = uploaded_file.name
                
                st.success("✅ Documentation generated successfully!")
                
            except json.JSONDecodeError as e:
                st.error(f"❌ Failed to parse JSON from model: {e}")
                st.error("Raw response:")
                st.text_area("Raw model response", doc_string, height=200)
            except Exception as e:
                st.error(f"❌ Error generating documentation: {e}")

# Display generated documentation if it exists in session state
if hasattr(st.session_state, 'generated_docs') and st.session_state.generated_docs:
    st.divider()
    st.subheader("📋 Generated Documentation")
    
    # Display the JSON in an expandable section
    with st.expander("View JSON Documentation", expanded=True):
        st.json(st.session_state.generated_docs)
    
    # Download button (always visible if docs are generated)
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("📥 Download DOCX", type="secondary"):
            try:
                docx_buffer = build_docx(st.session_state.generated_docs, st.session_state.generated_filename)
                st.download_button(
                    label="💾 Click to Download DOCX",
                    data=docx_buffer,
                    file_name=st.session_state.generated_filename.replace(".py", "_documentation.docx"),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.error(f"❌ Error creating DOCX: {e}")
    
    # Clear button to reset
    with col2:
        if st.button("🗑️ Clear Generated Documentation"):
            st.session_state.generated_docs = None
            st.session_state.generated_filename = None
            st.rerun()

else:
    if uploaded_file:
        st.info("👆 Click 'Generate Documentation' to process your file")
    else:
        st.info("👆 Upload a Python file to get started") 